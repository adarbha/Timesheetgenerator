import argparse
import shutil
import re
import json
import calendar
import pandas as pd
import numpy as np
from docx import Document
from pandas.tseries.offsets import CustomBusinessDay

PREFIX = r"Procal_Tech_Weekly_Timesheet_"
WORK_HOURS = 8

def date_to_string(dtindex):
    '''Formatter for displaying dates'''
    return "{}/{}/{}".format(dtindex.month,dtindex.day,dtindex.year)

def change_week_start_end_dates(doc,week):
    '''Change start and end dates within the document. Returns nothing'''
    para_to_change = doc.paragraphs[4] #This is static and never changes
    pattern = re.compile(r'(Duration: Week Start Date:)(.*)(Week End Date:)(.*)',re.IGNORECASE)
    search = pattern.search(para_to_change.text)
    if search:
        para_to_change.text = "{} {} \t {} {}".format(search.group(1),date_to_string(min(week)),search.group(3),date_to_string(max(week))) 
    return True   

def fill_work_days_table(doc,week):
    '''Fills the table with appropriate hours for the week'''
    first_day_of_week = min(week).weekday() + 1
    last_day_of_week = max(week).weekday() + 1
    rows = doc.tables[1].rows #This is static and should never change
    for day in range(1,first_day_of_week):
        cells = rows[day].cells
        for c in cells[2:]:
            c.text = "--"

    if last_day_of_week < 5:
        for day in range(last_day_of_week + 1,6):
            cells = rows[day].cells
            for c in cells[2:]:
                c.text = "--"

    return True

def fill_work_hours(doc,week):
    '''Fill work hours at the bottom for this week'''
    week_work_hours = len(week) * WORK_HOURS
    work_hour_cells = doc.tables[1].rows[8].cells #This is very static
    work_hour_cells[-1].text = str(week_work_hours)
    work_hour_cells[-7].text = str(week_work_hours)
    return True

def fill_dates_for_signature(doc):
    '''Fill dates for signature'''
    cells = doc.tables[1].rows[10].cells
    cells[5].text = date_to_string(pd.datetime.now())

# Command line arguments to point to a config. Config should contain all you need / or make it contain
parser = argparse.ArgumentParser(description="Create a bunch of timesheets b/n user set dates and save your time and world")
parser.add_argument('config',help="path to config.json")
args = parser.parse_args()

# Snoopin things from config
with open(args.config,'r') as f:
    config = json.loads(f.read())

#Building a business calendar for that month
bus_days = CustomBusinessDay(holidays=config['holidays'])
bus_dates = pd.date_range(start=config['start_date'],end=config['end_date'],freq=bus_days)
#Calender days by week
array_of_weeks = [bus_dates[bus_dates.week == w] for w in range(min(bus_dates.week),max(bus_dates.week) + 1)]

#Loop through array of weeks to create a time-sheet file for each week
for week in array_of_weeks:
    file_name = "{}{}_{}_{}_{}.docx".format(config['landing_dir'],PREFIX,calendar.month_abbr[week[0].month],min(week).day,max(week).day)
    shutil.copyfile(config['template_doc'],file_name)
    doc = Document(file_name)
    change_week_start_end_dates(doc,week)
    fill_work_days_table(doc,week)
    fill_work_hours(doc,week)
    fill_dates_for_signature(doc)
    doc.save(file_name)





